"""
Генерация docx из шаблонов: подстановка плейсхолдеров {{...}} из данных заказа.
Шаблоны в папке templates/ в корне проекта (PROJECT_CONTEXT, раздел 13).
Используется простая замена строк (шаблоны с пробелами в плейсхолдерах, напр. «ФИО продавец»).
"""
from datetime import date
from pathlib import Path
from io import BytesIO
from typing import Dict, Optional

from docx import Document

# Папка шаблонов: корень проекта / templates
_BASE = Path(__file__).resolve().parent.parent.parent.parent
TEMPLATES_DIR = _BASE / "templates"

# Маппинг: имя плейсхолдера в шаблоне (без {{ }}) → ключ в form_data
PLACEHOLDER_TO_FIELD = {
    "ФИО": "client_fio",
    "Паспорт": "client_passport",
    "Адрес": "client_address",
    "Телефон": "client_phone",
    "ФИО продавец": "seller_fio",
    "Паспорт продавец": "seller_passport",
    "Адрес продавец": "seller_address",
    "VIN": "vin",
    "Марка, модель": "brand_model",
    "Тип ТС": "vehicle_type",
    "Год выпуска": "year",
    "Двигатель": "engine",
    "№ шасси (рамы)": "chassis",
    "№ кузова": "body",
    "Цвет": "color",
    "СРТС": "srts",
    "Гос. Номер": "plate_number",
    "ПТС": "pts",
    "Сумма ДКП": "summa_dkp",
    "Дата ДКП": None,  # подставим дату отдельно
}


def _form_data_to_replace_map(form_data: Optional[dict], doc_date: Optional[date] = None) -> Dict[str, str]:
    """Словарь подстановки: «{{ ключ }}» → значение."""
    if not form_data:
        form_data = {}
    doc_date = doc_date or date.today()
    result = {}
    for placeholder, field_key in PLACEHOLDER_TO_FIELD.items():
        if field_key is None:
            result[placeholder] = doc_date.strftime("%d.%m.%Y")
            continue
        value = form_data.get(field_key)
        if value is None:
            value = ""
        result[placeholder] = str(value)
    return result


def _replace_in_paragraph(paragraph, replace_map: dict[str, str]) -> None:
    text = paragraph.text
    for key, value in replace_map.items():
        text = text.replace("{{" + key + "}}", value)
        text = text.replace("{{ " + key + " }}", value)
    if text != paragraph.text:
        paragraph.clear()
        paragraph.add_run(text)


def render_docx(template_name: str, form_data: Optional[dict], doc_date: Optional[date] = None) -> bytes:
    """
    Генерирует docx из шаблона (например DKP.docx), подставляя {{ плейсхолдер }} из form_data.
    Возвращает файл как bytes.
    """
    path = TEMPLATES_DIR / template_name
    if not path.is_file():
        raise FileNotFoundError(f"Шаблон не найден: {template_name}")
    doc = Document(str(path))
    replace_map = _form_data_to_replace_map(form_data, doc_date)
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replace_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replace_map)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
