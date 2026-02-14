"""
Однократный скрипт: добавить в шаблон zaiavlenie.docx плейсхолдеры для юр. лица:
{{Название}}, {{ИНН}}, {{ОГРН}} в блоке «СВЕДЕНИЯ О СОБСТВЕННИКЕ ТС».
Запуск из корня проекта: python scripts/patch_zaiavlenie_placeholders.py
"""
from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE / "templates" / "zaiavlenie.docx"


def set_cell_text(cell, text: str) -> None:
    """Установить текст ячейки."""
    for p in cell.paragraphs:
        p.clear()
        p.add_run(text)


def find_table_with_owner_section(doc: Document):
    """Найти таблицу с блоком «СВЕДЕНИЯ О СОБСТВЕННИКЕ ТС»."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "СВЕДЕНИЯ О СОБСТВЕННИКЕ" in cell.text:
                    return table
    return None


def main():
    doc = Document(str(TEMPLATE_PATH))
    table = find_table_with_owner_section(doc)
    if not table:
        print("Таблица «СВЕДЕНИЯ О СОБСТВЕННИКЕ» не найдена.")
        return

    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        first_text = cells[0].text.strip()
        # ИНН — во второй ячейке вставить {{ИНН}}
        if first_text == "ИНН" and len(cells) >= 2:
            set_cell_text(cells[1], "{{ИНН}}")
        # Ячейка с {{ФИО}} — добавить {{Название}} для юр. лица (подстановка по очереди)
        for cell in cells:
            if "{{ФИО}}" in cell.text:
                set_cell_text(cell, "{{ФИО}} {{Название}}")
                break
        # ОГРН — если есть строка с подписью ОГРН, заполнить значение
        if first_text == "ОГРН" and len(cells) >= 2:
            set_cell_text(cells[1], "{{ОГРН}}")

    # Добавить строку ОГРН в конец таблицы, если такой строки не было
    has_ogrn = any(
        "ОГРН" in row.cells[0].text
        for row in table.rows
        if row.cells
    )
    if not has_ogrn:
        new_row = table.add_row()
        if len(new_row.cells) >= 2:
            set_cell_text(new_row.cells[0], "ОГРН")
            set_cell_text(new_row.cells[1], "{{ОГРН}}")

    doc.save(str(TEMPLATE_PATH))
    print("Готово: в шаблон zaiavlenie.docx добавлены плейсхолдеры {{Название}}, {{ИНН}}, {{ОГРН}}.")


if __name__ == "__main__":
    main()
